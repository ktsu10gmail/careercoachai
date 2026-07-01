from pathlib import Path
import re

from fastapi import HTTPException, UploadFile, status

ALLOWED_RESUME_EXTENSIONS = {".txt", ".pdf", ".docx"}
PROJECT_ROOT = Path(__file__).resolve().parents[2]


def validate_resume_upload(resume: UploadFile) -> str:
    filename = resume.filename or ""
    extension = Path(filename).suffix.lower()
    if extension not in ALLOWED_RESUME_EXTENSIONS:
        allowed = ", ".join(sorted(ALLOWED_RESUME_EXTENSIONS))
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported resume format. Upload one of: {allowed}.",
        )
    return extension


def extract_resume_text(path: Path) -> str:
    path = resolve_resume_path(path)
    extension = path.suffix.lower()
    if extension == ".txt":
        return _normalize_extracted_text(path.read_text(encoding="utf-8", errors="ignore"))
    if extension == ".pdf":
        return _extract_pdf_text(path)
    if extension == ".docx":
        return _extract_docx_text(path)
    return ""


def resolve_resume_path(path: Path | str) -> Path:
    resume_path = Path(path)
    if resume_path.is_absolute():
        return resume_path

    project_relative_path = PROJECT_ROOT / resume_path
    if project_relative_path.exists():
        return project_relative_path

    return Path.cwd() / resume_path


def _extract_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF parsing dependency is not installed.") from exc

    reader = PdfReader(str(path))
    page_texts = []
    for page in reader.pages:
        text = ""
        try:
            text = page.extract_text(extraction_mode="layout") or ""
        except TypeError:
            text = page.extract_text() or ""
        except Exception:
            text = page.extract_text() or ""
        page_texts.append(text)
    return _normalize_extracted_text("\n\n".join(page_texts))


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX parsing dependency is not installed.") from exc

    document = Document(str(path))
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    for section in document.sections:
        parts.extend(paragraph.text for paragraph in section.header.paragraphs)
        parts.extend(paragraph.text for paragraph in section.footer.paragraphs)
    return _normalize_extracted_text("\n".join(parts))


def _normalize_extracted_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    text = re.sub(r"[ \t\r\f\v]+", " ", text)
    lines = [re.sub(r" +", " ", line).strip() for line in text.splitlines()]
    lines = [line for line in lines if line]
    return "\n".join(lines).strip()
